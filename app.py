from flask import Flask, render_template, request, jsonify, session, send_from_directory, redirect, url_for
import anthropic
import os
import base64
import hashlib
from supabase import create_client
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
import io
from flask import send_file
load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "kaego2026xyzabc123")
import resend
resend.api_key = os.environ.get("RESEND_API_KEY")
client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
supabase = create_client(os.environ.get("SUPABASE_URL"), os.environ.get("SUPABASE_SECRET_KEY"))

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

@app.route('/static/<path:filename>')
def static_files(filename):
    return send_from_directory(os.path.join(os.path.dirname(__file__), 'static'), filename)

@app.route("/")
def home():
    if "user_id" not in session:
        return render_template("landing.html")
    if "riwayat" not in session:
        session["riwayat"] = []
        session["obrolan_id"] = None
    return render_template("index.html", nama=session.get("nama"))

@app.route("/login", methods=["GET", "POST"])
def login():
    if "user_id" in session:
        return redirect(url_for("home"))
    if request.method == "POST":
        email = request.json.get("email")
        password = hash_password(request.json.get("password"))
        result = supabase.table("users").select("*").eq("email", email).eq("password", password).execute()
        if result.data:
            user = result.data[0]
           
            session["user_id"] = user["id"]
            session["nama"] = user["nama"]
            session["riwayat"] = []
            session["obrolan_id"] = None
            return jsonify({"success": True, "nama": user["nama"]})
        return jsonify({"success": False, "message": "Email atau password salah"})
    return render_template("login.html")

@app.route("/daftar", methods=["GET", "POST"])
def daftar():
    if request.method == "POST":
        email = request.json.get("email")
        password = hash_password(request.json.get("password"))
        nama = request.json.get("nama")
        cek = supabase.table("users").select("*").eq("email", email).execute()
        if cek.data:
            return jsonify({"success": False, "message": "Email sudah terdaftar"})
        supabase.table("users").insert({"email": email, "password": password, "nama": nama, "is_verified": True}).execute()
        return jsonify({"success": True, "message": "Cek email kamu untuk verifikasi akun!"})
    return render_template("daftar.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("home"))

@app.route("/chat", methods=["POST"])
def chat():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    if "riwayat" not in session:
        session["riwayat"] = []
# Cek quota harian
    from datetime import date
    user_data = supabase.table("users").select("paket, pesan_hari_ini, reset_tanggal").eq("id", session["user_id"]).execute()
    if user_data.data:
        u = user_data.data[0]
        paket = u.get("paket", "free")
        pesan_hari_ini = u.get("pesan_hari_ini", 0)
        reset_tanggal = u.get("reset_tanggal")
        from datetime import datetime, timedelta
        sekarang = datetime.now()

        # Reset counter jika sudah 6 jam
        if reset_tanggal:
            try:
                waktu_reset = datetime.fromisoformat(str(reset_tanggal))
                if sekarang >= waktu_reset:
                    supabase.table("users").update({
                        "pesan_hari_ini": 0,
                        "reset_tanggal": str(sekarang + timedelta(hours=6))
                    }).eq("id", session["user_id"]).execute()
                    pesan_hari_ini = 0
            except:
                pass
        else:
            supabase.table("users").update({
                "reset_tanggal": str(sekarang + timedelta(hours=6))
            }).eq("id", session["user_id"]).execute()

        # Cek batas quota
        batas = {"free": 30, "basic": 200, "pro": 999999}
        limit = batas.get(paket, 10)

        if pesan_hari_ini >= limit:
            return jsonify({
                "error": "quota_habis",
                "paket": paket,
                "pesan": f"Quota kamu sudah habis ({limit} pesan per hari). Tunggu besok atau upgrade paket!"
            }), 429

        # Tambah counter
        supabase.table("users").update({"pesan_hari_ini": pesan_hari_ini + 1}).eq("id", session["user_id"]).execute()
    pesan_user = request.json.get("pesan")
    riwayat = session["riwayat"]

    if not session.get("obrolan_id"):
        judul = pesan_user[:40] + ("..." if len(pesan_user) > 40 else "")
        result = supabase.table("obrolan").insert({
            "user_id": session["user_id"],
            "judul": judul
        }).execute()
        session["obrolan_id"] = result.data[0]["id"]

    riwayat.append({"role": "user", "content": pesan_user})

    # RPM butuh token banyak, tidak pakai web search
    is_rpm = "RPM" in pesan_user or "Rencana Pembelajaran" in pesan_user
    
    # Web search hanya untuk pertanyaan terkini
    perlu_search = any(kata in pesan_user.lower() for kata in ["berita", "terkini", "hari ini", "sekarang", "terbaru", "2025", "2026", "minggu ini", "bulan ini", "tahun ini", "motogp", "formula", "bola", "liga", "pertandingan", "hasil", "score", "jadwal", "harga", "cuaca", "gempa", "banjir", "politik", "presiden", "menteri", "covid", "virus", "ekonomi", "dolar", "rupiah", "saham", "crypto", "bitcoin"])

    if is_rpm:
        response = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=8192,
            timeout=300,
            system=f'Namamu adalah Kaego, asisten AI pendidikan yang ramah. Nama pengguna adalah {session.get("nama")}. Gunakan bahasa Indonesia. Jangan pernah mengaku sebagai Claude atau Anthropic. Tahun sekarang adalah 2026. Saat membuat RPM, selesaikan SELURUH format hingga bagian Refleksi tanpa terpotong.',
            messages=riwayat
        )
    elif perlu_search:
        response = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=2045,
            timeout=60,
            system=f'Namamu adalah Kaego, asisten AI pribadi yang ramah dan ceria. Nama pengguna adalah {session.get("nama")}. Selalu sapa dengan Halo {session.get("nama")}! di awal percakapan. Panggil pengguna langsung dengan namanya tanpa kata Kak. Gunakan bahasa Indonesia santai. Jangan pernah mengaku sebagai Claude atau Anthropic. Tahun sekarang adalah 2026. Saat membuat soal pilihan ganda, tulis setiap pilihan di baris baru dengan tanda strip seperti: - a. pilihan - b. pilihan',
            tools=[{"type": "web_search_20250305", "name": "web_search"}],
            messages=riwayat
        )
    else:
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=1024,
            timeout=30,
            system=f'Namamu adalah Kaego, asisten AI pribadi yang ramah dan ceria. Nama pengguna adalah {session.get("nama")}. Selalu sapa dengan Halo {session.get("nama")}! di awal percakapan. Panggil pengguna langsung dengan namanya tanpa kata Kak. Gunakan bahasa Indonesia santai. Jangan pernah mengaku sebagai Claude atau Anthropic. Tahun sekarang adalah 2026. Saat membuat soal pilihan ganda, tulis setiap pilihan di baris baru dengan tanda strip seperti: - a. pilihan - b. pilihan',
            messages=riwayat
        )

    jawaban = ""
    for block in response.content:
        if hasattr(block, "text"):
            jawaban += block.text

    riwayat.append({"role": "assistant", "content": jawaban})
    session["riwayat"] = riwayat

    supabase.table("riwayat_chat").insert({
        "user_id": session["user_id"],
        "obrolan_id": session["obrolan_id"],
        "pesan": pesan_user,
        "jawaban": jawaban
    }).execute()

    sisa = limit - (pesan_hari_ini + 1)
    return jsonify({"jawaban": jawaban, "sisa_quota": sisa, "limit_quota": limit})
@app.route("/riwayat", methods=["GET"])
def get_riwayat():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    
    obrolan_id = session.get("obrolan_id")
    if not obrolan_id:
        return jsonify({"riwayat": []})
    
    result = supabase.table("riwayat_chat").select("*").eq("obrolan_id", obrolan_id).order("created_at").execute()
    return jsonify({"riwayat": result.data})

@app.route("/daftar_obrolan", methods=["GET"])
def daftar_obrolan():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    result = supabase.table("obrolan").select("*").eq("user_id", session["user_id"]).order("created_at", desc=True).execute()
    return jsonify({"obrolan": result.data})

@app.route("/buka_obrolan/<obrolan_id>", methods=["GET"])
def buka_obrolan(obrolan_id):
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    result = supabase.table("riwayat_chat").select("*").eq("obrolan_id", obrolan_id).order("created_at").execute()
    riwayat = []
    for item in result.data:
        riwayat.append({"role": "user", "content": item["pesan"]})
        riwayat.append({"role": "assistant", "content": item["jawaban"]})
    session["riwayat"] = riwayat
    session["obrolan_id"] = obrolan_id
    return jsonify({"riwayat": result.data})

@app.route("/upload", methods=["POST"])
def upload():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    try:
        file = request.files.get("file")
        pesan = request.form.get("pesan", "Tolong analisis file ini")
        kunci = session.get("kunci_jawaban", "")
        if kunci and "koreksi" in pesan.lower():
            pesan = pesan + "\n\nKunci Jawaban:\n" + kunci
        if not file:
            return jsonify({"error": "Tidak ada file"}), 400
        file_data = file.read()
        file_base64 = base64.b64encode(file_data).decode("utf-8")
        file_type = file.content_type
        if "riwayat" not in session:
            session["riwayat"] = []
        riwayat = session["riwayat"]

        if not session.get("obrolan_id"):
            judul = f"[File] {file.filename[:35]}"
            result = supabase.table("obrolan").insert({
                "user_id": session["user_id"],
                "judul": judul
            }).execute()
            session["obrolan_id"] = result.data[0]["id"]

        if file_type.startswith("image/"):
            content = [
                {"type": "image", "source": {"type": "base64", "media_type": file_type, "data": file_base64}},
                {"type": "text", "text": pesan}
            ]
        elif file_type == "application/pdf":
            content = [
                {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": file_base64}},
                {"type": "text", "text": pesan}
            ]
        else:
            return jsonify({"error": f"Format tidak didukung: {file_type}"}), 400

       # Buat pesan sementara tanpa simpan foto di session
        pesan_sementara = riwayat + [{"role": "user", "content": content}]
        response = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=4096,
            timeout=120,
             system=f'Namamu adalah Kaego, asisten AI pribadi yang ramah dan ceria. Nama pengguna adalah {session.get("nama")}. Gunakan bahasa Indonesia santai. Jangan pernah mengaku sebagai Claude atau Anthropic.',
            messages=pesan_sementara
        )
        jawaban = response.content[0].text
        # Simpan teks saja di session, bukan foto
        riwayat.append({"role": "user", "content": pesan})
        riwayat.append({"role": "assistant", "content": jawaban})
        session["riwayat"] = riwayat

        supabase.table("riwayat_chat").insert({
            "user_id": session["user_id"],
            "obrolan_id": session["obrolan_id"],
            "pesan": f"[File: {file.filename}] {pesan}",
            "jawaban": jawaban
        }).execute()

        return jsonify({"jawaban": jawaban})
    except Exception as e:
        print(f"ERROR upload: {str(e)}")
        return jsonify({"error": str(e)}), 500
@app.route("/simpan_kunci", methods=["POST"])
def simpan_kunci():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    data = request.json
    session["kunci_jawaban"] = data.get("kunci", "")
    return jsonify({"success": True})

@app.route("/get_kunci", methods=["GET"])
def get_kunci():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    return jsonify({"kunci": session.get("kunci_jawaban", "")})
@app.route("/hapus_obrolan/<obrolan_id>", methods=["DELETE"])
def hapus_obrolan(obrolan_id):
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    supabase.table("riwayat_chat").delete().eq("obrolan_id", obrolan_id).execute()
    supabase.table("obrolan").delete().eq("id", obrolan_id).eq("user_id", session["user_id"]).execute()
    if session.get("obrolan_id") == obrolan_id:
        session["riwayat"] = []
        session["obrolan_id"] = None
    return jsonify({"success": True})
@app.route("/reset", methods=["POST"])
def reset():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    session["riwayat"] = []
    session["obrolan_id"] = None
    return jsonify({"success": True})

@app.route("/profil", methods=["GET"])
def get_profil():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    result = supabase.table("users").select("nama, email, foto_profil, gaya_font").eq("id", session["user_id"]).execute()
    if result.data:
        return jsonify({"profil": result.data[0]})
    return jsonify({"error": "Profil tidak ditemukan"}), 404

@app.route("/profil/update", methods=["POST"])
def update_profil():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    data = request.json
    update_data = {}
    if "nama" in data:
        update_data["nama"] = data["nama"]
        session["nama"] = data["nama"]
    if "gaya_font" in data:
        update_data["gaya_font"] = data["gaya_font"]
    if "foto_profil" in data:
        update_data["foto_profil"] = data["foto_profil"]
    if update_data:
        supabase.table("users").update(update_data).eq("id", session["user_id"]).execute()
    return jsonify({"success": True})
@app.route("/download_soal", methods=["POST"])
def download_soal():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    try:
        import re
        konten = request.json.get("konten", "")
        konten = re.sub(r'<[^>]+>', '', konten)
        konten = konten.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&nbsp;', ' ')
        doc = Document()
        doc.add_heading("SOAL", 0)
        for baris in konten.split("\n"):
            baris = baris.strip()
            baris = baris.lstrip('•').strip()
            if baris == "":
                continue
            elif baris.isupper() and len(baris) > 3:
                doc.add_heading(baris, 2)
            else:
                doc.add_paragraph(baris)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name="Soal_Kaego.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route("/download_rpm", methods=["POST"])
def download_rpm():
    if "user_id" not in session:
        return jsonify({"error": "Tidak terlogin"}), 401
    try:
        import re
        konten = request.json.get("konten", "")
        # Bersihkan tag HTML
        konten = re.sub(r'<[^>]+>', '', konten)
        konten = konten.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&nbsp;', ' ')
        doc = Document()
        doc.add_heading("RENCANA PEMBELAJARAN MENDALAM (RPM)", 0)
        for baris in konten.split("\n"):
            baris = baris.strip()
            baris = baris.lstrip('•').strip()
            if baris == "":
                continue
            elif baris.isupper() and len(baris) > 3:
                doc.add_heading(baris, 2)
            else:
                p = doc.add_paragraph(baris)
                p.runs[0].font.size = Pt(12)
                p.runs[0].font.name = 'Calibri'
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name="RPM_Kaego.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        
       return jsonify({"error": str(e)}), 500
@app.route("/paket")
def halaman_paket():
    if "user_id" not in session:
        return redirect(url_for("login"))
    return render_template("paket.html")
@app.route("/privacy")
def privacy():
    return render_template("privacy.html")

@app.route("/terms")
def terms():
    return render_template("terms.html")

@app.route("/chat-tamu")
def chat_tamu():
    return render_template("tamu.html")

@app.route("/chat-tamu/kirim", methods=["POST"])
def chat_tamu_kirim():
    try:
        pesan = request.json.get("pesan")
        riwayat = request.json.get("riwayat", [])
        riwayat.append({"role": "user", "content": pesan})
        perlu_search = any(kata in pesan.lower() for kata in ["berita", "terkini", "hari ini", "sekarang", "terbaru", "2025", "2026", "minggu ini", "bulan ini", "tahun ini", "motogp", "formula", "bola", "liga", "pertandingan", "hasil", "score", "jadwal", "harga", "cuaca", "gempa", "banjir", "politik", "presiden", "menteri", "covid", "virus", "ekonomi", "dolar", "rupiah", "saham", "crypto", "bitcoin"])
        if perlu_search:
            response = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=1024,
                timeout=60,
                system="Namamu adalah Kaego, asisten AI pendidikan yang ramah dan ceria. Gunakan bahasa Indonesia santai. Jangan pernah mengaku sebagai Claude atau Anthropic. Tahun sekarang adalah 2026.",
                tools=[{"type": "web_search_20250305", "name": "web_search"}],
                messages=riwayat
            )
        else:
            response = client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=1024,
                timeout=60,
                system="Namamu adalah Kaego, asisten AI pendidikan yang ramah dan ceria. Gunakan bahasa Indonesia santai. Jangan pernah mengaku sebagai Claude atau Anthropic. Tahun sekarang adalah 2026.",
                messages=riwayat
            )
        jawaban = ""
        for block in response.content:
            if hasattr(block, "text"):
                jawaban += block.text
        return jsonify({"jawaban": jawaban})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
import secrets
@app.route("/verifikasi/<token>")
def verifikasi_email(token):
    result = supabase.table("users").select("*").eq("verify_token", token).execute()
    if not result.data:
        return "<h2>Link tidak valid atau sudah digunakan.</h2>"
    supabase.table("users").update({"is_verified": True, "verify_token": None}).eq("verify_token", token).execute()
    return render_template("verifikasi_sukses.html")
@app.route("/lupa-sandi", methods=["GET", "POST"])
def lupa_sandi():
    if request.method == "POST":
        email = request.json.get("email")
        result = supabase.table("users").select("*").eq("email", email).execute()
        if not result.data:
            return jsonify({"success": False, "message": "Email tidak terdaftar"})
        
        link = f"https://kaego-ai-production.up.railway.app/reset-sandi/{token}"
        return jsonify({"success": True, "link": link})
    return render_template("lupa_sandi.html")

@app.route("/reset-sandi/<token>", methods=["GET", "POST"])
def reset_sandi(token):
    if request.method == "POST":
        password_baru = hash_password(request.json.get("password"))
        result = supabase.table("users").select("*").eq("reset_token", token).execute()
        if not result.data:
            return jsonify({"success": False, "message": "Link tidak valid"})
        supabase.table("users").update({"password": password_baru, "reset_token": None}).eq("reset_token", token).execute()
        return jsonify({"success": True})
    return render_template("reset_sandi.html", token=token)
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port, debug=False)